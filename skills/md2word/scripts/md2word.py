#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown到Word文档转换工具
使用配置系统驱动格式化，支持自定义YAML配置和预设格式

配置说明详见: assets/presets/*.yaml 和 references/config-reference.md
"""

import os
import argparse
import re
import glob
import tempfile
import urllib.request
import urllib.parse
import io

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.shared import OxmlElement
from PIL import Image
from bs4 import BeautifulSoup

# 导入配置模块
from config import Config, load_config, get_preset, get_default_preset, list_presets, get_config, set_config

# 导入功能模块
from formatter import (
    convert_quotes_to_chinese,
    parse_text_formatting,
    set_run_format_with_styles,
    set_paragraph_format,
    parse_alignment,
    extract_alignment,
    hex_to_rgb,
)
from table_handler import (
    is_table_row,
    create_word_table,
    create_word_table_from_html,
)
from chart_handler import create_mermaid_chart
from svg_handler import render_inline_svg
from footnote_handler import (
    FootnoteManager, extract_footnote_defs, NOTE_REF_RE,
    set_footnote_restart_per_section,
)


# ----------------------------------------------------------------------------
# 脚注/尾注：当前转换的 FootnoteManager（全局，供 parse_text_with_footnotes 使用）
# ----------------------------------------------------------------------------
_active_fn_manager = None


def parse_text_with_footnotes(paragraph, text, title_level=0, is_quote=False):
    """解析文本，识别行内 [^id] 脚注引用；其余走 parse_text_formatting。

    无激活的 FootnoteManager 时（脚注未启用），fallback 到 parse_text_formatting。
    """
    if _active_fn_manager is None:
        parse_text_formatting(paragraph, text, title_level=title_level, is_quote=is_quote)
        return
    # 遍历 [^id] 引用：引用之间的普通文本走 parse_text_formatting，引用处插入脚注引用 run
    last = 0
    for m in NOTE_REF_RE.finditer(text):
        if m.start() > last:
            parse_text_formatting(paragraph, text[last:m.start()],
                                  title_level=title_level, is_quote=is_quote)
        _active_fn_manager.add_reference(paragraph, m.group(1))
        last = m.end()
    if last < len(text):
        parse_text_formatting(paragraph, text[last:],
                              title_level=title_level, is_quote=is_quote)


# ============================================================================
# 图片处理
# ============================================================================

def get_image_output_path(md_file_path, png_filename):
    """获取图片输出路径，确保目录存在"""
    md_dir = os.path.dirname(os.path.abspath(md_file_path))
    md_filename_base = os.path.splitext(os.path.basename(md_file_path))[0]
    image_dir = os.path.join(md_dir, f"{md_filename_base}_images")

    if not os.path.exists(image_dir):
        try:
            os.makedirs(image_dir)
            print(f"📂 创建图片目录: {os.path.relpath(image_dir)}")
        except OSError as e:
            print(f"⚠️ 创建目录失败: {e}")
            return None

    return os.path.join(image_dir, png_filename)


def _postprocess_image_for_word(image, target_display_cm, target_dpi=260):
    """根据目标显示宽度与DPI对图像进行高质量下采样"""
    try:
        target_inches = float(target_display_cm) / 2.54
        target_px_width = max(1, int(target_inches * target_dpi))
        if image.width > target_px_width:
            new_height = int(image.height * (target_px_width / image.width))
            image = image.resize((target_px_width, new_height), Image.LANCZOS)
    except Exception:
        pass
    return image


def insert_image_to_word(doc, image):
    """将PIL图片对象插入Word文档"""
    config = get_config()
    image_config = config.get('image', {})
    page_config = config.get('page', {})

    display_ratio = image_config.get('display_ratio', 0.92)
    max_width_cm = image_config.get('max_width_cm', 14.2)
    target_dpi = image_config.get('target_dpi', 260)

    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
        page_width = page_config.get('width', 21.0)
        margin_left = page_config.get('margin_left', 3.18)
        margin_right = page_config.get('margin_right', 3.18)
        available_width_cm = page_width - margin_left - margin_right
        target_display_cm = min(available_width_cm * display_ratio, max_width_cm)
        image = _postprocess_image_for_word(image, target_display_cm, target_dpi=target_dpi)
        try:
            image.save(temp_file.name, format='PNG', optimize=True, compress_level=9)
        except Exception:
            image.save(temp_file.name, format='PNG', optimize=True)
        temp_filename = temp_file.name

    try:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        img_width_cm = target_display_cm
        run = paragraph.add_run()
        run.add_picture(temp_filename, width=Cm(img_width_cm))
    finally:
        try:
            os.unlink(temp_filename)
        except:
            pass


def download_external_image(url):
    """从URL下载图片并返回PIL Image对象"""
    try:
        req = urllib.request.Request(url, headers={
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36'
        })
        with urllib.request.urlopen(req, timeout=20) as response:
            image_data = response.read()
        image = Image.open(io.BytesIO(image_data))
        image.load()  # 确保数据已加载
        return image
    except Exception as e:
        print(f"⚠️  图片下载失败: {url[:80]}... ({e})")
        return None


# ============================================================================
# 文档结构元素
# ============================================================================

def add_horizontal_line(doc):
    """添加分割线"""
    config = get_config()
    hr_config = config.get('horizontal_rule', {})
    
    p = doc.add_paragraph()
    p.alignment = parse_alignment(hr_config.get('alignment', 'center'))
    
    character = hr_config.get('character', '─')
    repeat_count = hr_config.get('repeat_count', 55)
    run = p.add_run(character * repeat_count)
    
    font_name = hr_config.get('font', 'Times New Roman')
    font_size = hr_config.get('size', 12)
    color_hex = hr_config.get('color', '#808080')
    
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = hex_to_rgb(color_hex)


def add_task_list(doc, line):
    """添加任务列表"""
    config = get_config()
    task_config = config.get('lists.task', {})
    
    is_checked = line.startswith(('- [x]', '- [X]'))
    text = line[5:].strip()
    p = doc.add_paragraph()
    
    checked_mark = task_config.get('checked', '☑')
    unchecked_mark = task_config.get('unchecked', '☐')
    checkbox_run = p.add_run(f'{checked_mark} ' if is_checked else f'{unchecked_mark} ')
    set_run_format_with_styles(checkbox_run, {}, title_level=0)
    parse_text_formatting(p, text)
    set_paragraph_format(p)


def add_bullet_list(doc, line):
    """添加无序列表"""
    config = get_config()
    bullet_config = config.get('lists.bullet', {})
    
    text = line[2:].strip()
    p = doc.add_paragraph()
    
    marker = bullet_config.get('marker', '•')
    bullet_run = p.add_run(f'{marker} ')
    set_run_format_with_styles(bullet_run, {}, title_level=0)
    parse_text_formatting(p, text)
    set_paragraph_format(p)


def add_numbered_list(doc, line):
    """添加有序列表"""
    p = doc.add_paragraph()
    parse_text_formatting(p, line)
    set_paragraph_format(p)


def add_quote(doc, text):
    """添加引用块"""
    config = get_config()
    quote_config = config.get('quote', {})
    
    lines = text.split('\n')

    # 引用块不施加视觉样式（无底纹/无边框/无缩进），与正文一致
    bg_color = quote_config.get('background_color')       # None = 不加
    border_color = quote_config.get('border_color')       # None = 不加
    border_size = quote_config.get('border_size', 0)
    left_indent = quote_config.get('left_indent_inches', 0)
    font_size = quote_config.get('font_size')             # None = 继承正文
    line_spacing = quote_config.get('line_spacing')       # None = 继承正文

    for line_index, line in enumerate(lines):
        if not line.strip():
            p = doc.add_paragraph()
            if bg_color:
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), bg_color.lstrip('#'))
                pPr.append(shd)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            continue

        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        if bg_color:
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), bg_color.lstrip('#'))
            pPr.append(shd)
        if line_spacing:
            p.paragraph_format.line_spacing = line_spacing
        
        bullet_match = re.match(r'^\s*([-*+])\s+', line)
        number_match = re.match(r'^\s*(\d+\.)\s+', line)
        
        list_marker_run = None
        
        if bullet_match:
            indent_and_bullet = '    •  '
            list_marker_run = p.add_run(indent_and_bullet)
            line = line[bullet_match.end():]
        elif number_match:
            indent_and_number = f'    {number_match.group(1)} '
            list_marker_run = p.add_run(indent_and_number)
            line = line[number_match.end():]
        
        if list_marker_run and font_size:
            list_marker_run.font.size = Pt(font_size)
            set_run_format_with_styles(list_marker_run, {}, is_quote=True)

        parse_text_formatting(p, line, is_quote=True)
        set_paragraph_format(p, is_quote=True)

        if font_size:
            for run in p.runs:
                run.font.size = Pt(font_size)


def add_code_block(doc, code_lines, language):
    """添加代码块（出版级：等宽字体 + 浅灰底纹 + 细边框 + 关拼写检查）

    通过 code_block.content 配置控制：
    - font / east_asia_font / size / color：等宽字体与字号颜色
    - background_color：浅灰底纹（留空则不加底纹，向后兼容）
    - border_color / border_size：段落边框，相邻代码行边框在 Word 中自动连成完整框
    - no_proofread：关闭拼写检查（代码不应被拼写纠正）
    - left_indent / line_spacing：缩进与行距
    未配置 background_color / border 时行为同旧版。
    """
    config = get_config()
    code_config = config.get('code_block', {})

    label_config = code_config.get('label', {})
    label_enabled = label_config.get('enabled', True)
    if language and label_enabled:
        lang_p = doc.add_paragraph()
        lang_run = lang_p.add_run(f"[{language}]")
        lang_run.font.name = label_config.get('font', 'Times New Roman')
        lang_run.font.size = Pt(label_config.get('size', 10))
        lang_run.font.color.rgb = hex_to_rgb(label_config.get('color', '#808080'))

    content_config = code_config.get('content', {})
    left_indent = content_config.get('left_indent', 24)
    line_spacing = content_config.get('line_spacing', 1.2)
    font_name = content_config.get('font', 'Consolas')
    east_asia_font = content_config.get('east_asia_font', '等线')
    font_size = content_config.get('size', 9)
    color_hex = content_config.get('color', '#333333')
    bg_color = content_config.get('background_color')    # None → 不加底纹
    border_color = content_config.get('border_color')    # None → 不加边框
    border_size = content_config.get('border_size', 4)   # 1/8 pt 单位
    no_proof = content_config.get('no_proofread', True)

    for code_line in code_lines:
        p = doc.add_paragraph()
        run = p.add_run(code_line if code_line else ' ')
        # 等宽字体（ASCII/CS 用 font_name，东亚字符用 east_asia_font）
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.color.rgb = hex_to_rgb(color_hex)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:eastAsia'), east_asia_font)
        rFonts.set(qn('w:cs'), font_name)
        # 关闭拼写检查（代码不应被 Word 拼写纠正打扰）
        if no_proof:
            rPr.append(OxmlElement('w:noProof'))

        # 段落级底纹 + 边框：必须先于 spacing/ind append（CT_PPr 顺序 pBdr→shd→…→spacing→ind）
        pPr = p._p.get_or_add_pPr()
        if border_color:
            pBdr = OxmlElement('w:pBdr')
            for edge in ('top', 'left', 'bottom', 'right'):
                b = OxmlElement('w:' + edge)
                b.set(qn('w:val'), 'single')
                b.set(qn('w:sz'), str(border_size))
                b.set(qn('w:space'), '4')
                b.set(qn('w:color'), border_color.lstrip('#'))
                pBdr.append(b)
            pPr.append(pBdr)
        if bg_color:
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), bg_color.lstrip('#'))
            pPr.append(shd)

        # 段落格式：左缩进 + 行距 + 无首行缩进 + 无段前段后（python-docx 会插到 pBdr/shd 之后）
        pf = p.paragraph_format
        pf.left_indent = Pt(left_indent)
        pf.line_spacing = line_spacing
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = Pt(0)


def add_page_number(doc):
    """添加页码"""
    config = get_config()
    page_number_config = config.get('page_number', {})
    
    if not page_number_config.get('enabled', True):
        return
    
    try:
        section = doc.sections[0]
        footer = section.footer
        
        for para in footer.paragraphs:
            para.clear()
        
        if not footer.paragraphs:
            footer_para = footer.add_paragraph()
        else:
            footer_para = footer.paragraphs[0]
        
        position = page_number_config.get('position', 'center')
        if position == 'left':
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif position == 'right':
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        page_format = page_number_config.get('format', '1/x')
        if '1' in page_format:
            run = footer_para.add_run()
            fld_char_begin = parse_xml(r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            run._r.append(fld_char_begin)
            instr_text = parse_xml(r'<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGE </w:instrText>')
            run._r.append(instr_text)
            fld_char_end = parse_xml(r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            run._r.append(fld_char_end)
        
        if '/' in page_format:
            sep_run = footer_para.add_run("/")
        
        if 'x' in page_format:
            total_run = footer_para.add_run()
            fld_char_begin2 = parse_xml(r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            total_run._r.append(fld_char_begin2)
            instr_text2 = parse_xml(r'<w:instrText xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> NUMPAGES </w:instrText>')
            total_run._r.append(instr_text2)
            fld_char_end2 = parse_xml(r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            total_run._r.append(fld_char_end2)
        
        font_name = page_number_config.get('font', 'Times New Roman')
        font_size = page_number_config.get('size', 10.5)
        
        for run in footer_para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
            run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
    
    except Exception as e:
        print(f"⚠️  页码添加失败，将跳过页码设置: {e}")
        pass


# ============================================================================
# 工具函数
# ============================================================================

def find_template_file(auto: bool = False):
    """查找 `assets/templates/` 下的 .docx 模板。

    参数:
        auto: 默认 False → 返回 None（不自动加载模板，避免律所 logo 等视觉元素
              出现在用户没显式要求的 docx 里）。CLI 用 `--auto-template` 显式开启
              时传 True，才会真正去扫描 templates 目录。

    返回:
        模板文件绝对路径；找不到时返回 None。
    """
    if not auto:
        return None
    script_dir = os.path.dirname(os.path.abspath(__file__))
    skill_dir = os.path.dirname(script_dir)
    templates_dir = os.path.join(skill_dir, 'assets', 'templates')
    docx_files = glob.glob(os.path.join(templates_dir, "*.docx"))

    for docx_file in docx_files:
        filename = os.path.basename(docx_file).lower()
        if not any(keyword in filename for keyword in ['完整版', 'test', 'output', '输出']):
            if '模板' in filename or 'template' in filename:
                return docx_file

    return docx_files[0] if docx_files else None


def find_md_files():
    """查找脚本所在目录下的所有 .md 文件"""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_files = glob.glob(os.path.join(script_dir, "*.md"))
    return md_files


def generate_output_filename(md_file):
    """根据输入文件名生成输出文件名"""
    base_name = os.path.splitext(md_file)[0]
    return f"{base_name}.docx"


def debug_quotes_in_file(file_path):
    """简化的引号调试"""
    print("🔍 检查文件中的引号...")
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    ascii_double = content.count('"')
    chinese_open = content.count('"')
    chinese_close = content.count('"')
    
    print(f"📊 引号统计: ASCII双引号={ascii_double}, 中文开引号={chinese_open}, 中文闭引号={chinese_close}")
    
    for i, line in enumerate(content.split('\n'), 1):
        if '"' in line:
            print(f"🎯 测试第{i}行: {line.strip()}")
            _ = convert_quotes_to_chinese(line.strip())
            break
    
    print("-" * 30)


# ============================================================================
# 全书合并工具（--book 模式）
# ============================================================================

def rename_footnote_ids(content, ch):
    """全书合并预处理：给脚注 id 加章节前缀，避免跨章 [^id] 冲突。
    [^1] → [^1-1]（第 1 章的 1）；[^note] → [^2-note]（第 2 章的 note）。
    """
    return re.sub(r'\[\^([^\]]+)\]', lambda m: '[^%d-%s]' % (ch, m.group(1)), content)


def add_toc(doc):
    """在文档当前位置插入 Word 目录域（TOC field），打开后在 Word 中按 F9 更新。"""
    p = doc.add_paragraph()
    r1 = p.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin'); r1._r.append(b)
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = r'TOC \o "1-3" \h \z \u'; r1._r.append(instr)
    s = OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'), 'separate'); r1._r.append(s)
    placeholder = p.add_run('（目录：在 Word 中选中此处按 F9 更新）')
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end'); placeholder._r.append(e)


def add_book_header(section, title):
    """给 section 的页眉加书名（居中，小字）。"""
    if not title:
        return
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = hp.add_run(title)
    r.font.size = Pt(9)


def create_book(md_files, output_path, config, notes_mode='footnote'):
    """全书合并：多章 md → 单 docx。
    预处理：脚注 id 加章前缀防冲突；章间用 '---' 分隔（book_mode 下每个 --- 触发分页）。
    """
    print(f"📚 全书合并 {len(md_files)} 个文件 → {output_path}")
    merged = []
    for ch_idx, f in enumerate(md_files, 1):
        if not os.path.exists(f):
            print(f"⚠️  跳过不存在的文件: {f}")
            continue
        try:
            with open(f, 'r', encoding='utf-8') as fh:
                content = fh.read()
        except UnicodeDecodeError:
            with open(f, 'r', encoding='gbk') as fh:
                content = fh.read()
        content = rename_footnote_ids(content, ch_idx)
        merged.append(content)
        print(f"  第 {ch_idx} 章: {os.path.basename(f)}")
    if not merged:
        print("❌ 无可合并的章节")
        return
    full = '\n\n---\n\n'.join(merged)
    tmp_md = output_path + '.merged.md'
    with open(tmp_md, 'w', encoding='utf-8') as fh:
        fh.write(full)
    try:
        create_word_document(tmp_md, output_path, None, config, notes_mode, book_mode=True)
    finally:
        try:
            os.unlink(tmp_md)
        except OSError:
            pass


# ============================================================================
# 核心转换流程
# ============================================================================

def create_word_document(md_file_path, output_path, template_file=None, config: Config = None, notes_mode='footnote', book_mode=False):
    """从Markdown文件创建格式化的Word文档"""
    if config is None:
        config = get_config()

    print(f"📄 正在处理: {md_file_path}")
    print(f"📋 使用配置: {config.name}")

    if config.get('quotes.convert_to_chinese', True):
        debug_quotes_in_file(md_file_path)

    # 用于存储模板的header/footer XML元素
    template_header_xmls = []
    template_footer_xmls = []
    use_template_headers = False

    # 创建或加载文档
    if template_file and template_file != "none" and os.path.exists(template_file):
        print(f"📋 使用模板文件: {os.path.basename(template_file)}")

        # 直接使用模板文件，保留所有格式（包括页眉页脚）
        print("📄 直接打开模板文件")
        doc = Document(template_file)

        # 清空模板中的正文内容（保留页眉页脚和sectPr）
        # 获取body元素
        body = doc._element.body

        # 记住sectPr的位置和内容
        sectPr = body.find(qn('w:sectPr'))

        # 移除body中的所有子元素（除了sectPr）
        for child in list(body):
            if child.tag != qn('w:sectPr'):
                body.remove(child)

        use_template_headers = True
        print("✅ 已清空模板内容，保留页眉页脚")
    else:
        print("📄 创建新文档（不使用模板）")
        doc = Document()
        template_header_xmls = []
        template_footer_xmls = []
        template_header_rels = []
        template_footer_rels = []
        template_media_files = []
        template_sectPr_refs = []
        template_doc_rels = {}

    # 设置默认字体
    try:
        normal_style = doc.styles['Normal']
        font_config = config.get('fonts.default', {})
        normal_style.font.name = font_config.get('ascii', 'Times New Roman')
        normal_style.font.size = Pt(font_config.get('size', 10.5))
        normal_style._element.rPr.rFonts.set(qn('w:ascii'), font_config.get('ascii', 'Times New Roman'))
        normal_style._element.rPr.rFonts.set(qn('w:hAnsi'), font_config.get('ascii', 'Times New Roman'))
        normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), font_config.get('name', '仿宋_GB2312'))
        normal_style._element.rPr.rFonts.set(qn('w:cs'), font_config.get('ascii', 'Times New Roman'))
    except Exception as _:
        pass

    # 设置页面大小和页边距
    for section in doc.sections:
        page_config = config.get('page', {})
        orientation = page_config.get('orientation', 'portrait')
        if orientation == 'landscape':
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Cm(page_config.get('height', 29.7))
            section.page_height = Cm(page_config.get('width', 21.0))
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Cm(page_config.get('width', 21.0))
            section.page_height = Cm(page_config.get('height', 29.7))
        section.top_margin = Cm(page_config.get('margin_top', 2.54))
        section.bottom_margin = Cm(page_config.get('margin_bottom', 2.54))
        section.left_margin = Cm(page_config.get('margin_left', 3.18))
        section.right_margin = Cm(page_config.get('margin_right', 3.18))
        if book_mode:
            add_book_header(section, config.get('book.title', ''))
    
    # 读取Markdown文件
    try:
        with open(md_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except UnicodeDecodeError:
        with open(md_file_path, 'r', encoding='gbk') as f:
            content = f.read()

    # 去除 HTML 注释，避免渲染到 Word 文档中
    content = re.sub(r'<!--.*?-->', '', content, flags=re.DOTALL)

    lines = content.split('\n')
    # 脚注/尾注：提取 [^id]: 定义行（从正文移除），建立 FootnoteManager
    global _active_fn_manager
    fn_defs, lines = extract_footnote_defs(lines)
    fn_manager = FootnoteManager(notes_mode)
    fn_manager.set_defs(fn_defs)
    _active_fn_manager = fn_manager
    has_body_before_first_h2 = False
    has_seen_h2 = False
    has_seen_first_hr = False  # 追踪第一个分隔符
    i = 0
    svg_counter = [0]  # 内联 SVG 计数（用于命名输出文件）

    # 全书合并模式：正文前插入目录域 + 分页
    if book_mode:
        add_toc(doc)
        doc.add_page_break()

    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Mermaid 图表
        if re.match(r'^```\s*mermaid\b', line):
            mermaid_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                mermaid_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            
            if mermaid_lines:
                mermaid_code = '\n'.join(mermaid_lines)
                create_mermaid_chart(
                    doc,
                    lambda img: insert_image_to_word(doc, img),
                    get_image_output_path,
                    lambda: doc.add_paragraph(),
                    lambda p: set_paragraph_format(p),
                    mermaid_code,
                    md_file_path
                )
                if not has_seen_h2:
                    has_body_before_first_h2 = True
                print(f"✅ 处理Mermaid图表")
            continue
        
        # 内联 SVG 块（<svg>...</svg>，渲染为 PNG 嵌入；失败降级代码框显示源码）
        if line.startswith('<svg'):
            svg_lines = [lines[i]]
            i += 1
            while i < len(lines) and '</svg>' not in svg_lines[-1]:
                svg_lines.append(lines[i])
                i += 1
            svg_code = '\n'.join(svg_lines)
            ok = render_inline_svg(
                lambda img: insert_image_to_word(doc, img),
                svg_code, md_file_path, svg_counter[0])
            svg_counter[0] += 1
            if not ok:
                # 降级：把 SVG 源码当代码框显示
                add_code_block(doc, svg_code.split('\n'), 'svg')
            if not has_seen_h2:
                has_body_before_first_h2 = True
            continue

        # 代码块
        if line.startswith('```'):
            code_lines = []
            language = line[3:].strip()
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            add_code_block(doc, code_lines, language)
            if not has_seen_h2:
                has_body_before_first_h2 = True
            print("✅ 处理代码块")
            continue
        
        # HTML 表格
        if '<table>' in line.lower():
            html_table_content = []
            while i < len(lines):
                html_table_content.append(lines[i])
                if '</table>' in lines[i].lower():
                    i += 1
                    break
                i += 1
            if html_table_content:
                create_word_table_from_html(doc, '\n'.join(html_table_content), md_file_path)
                if not has_seen_h2:
                    has_body_before_first_h2 = True
            continue

        # HTML 块级元素（<div>, <p>, <span>, <section>, <article> 等）
        html_block_match = re.match(r'^<(div|p|span|section|article)\b([^>]*)>', line, re.IGNORECASE)
        if html_block_match:
            tag_name = html_block_match.group(1).lower()
            close_tag = f'</{tag_name}>'
            style_attr = html_block_match.group(2)
            html_lines = [line]
            # 如果闭合标签在同一行
            if close_tag in line.lower():
                pass
            else:
                while i + 1 < len(lines):
                    i += 1
                    html_lines.append(lines[i])
                    if close_tag in lines[i].lower():
                        break
            # 提取内容
            block_html = '\n'.join(html_lines)
            soup_html = BeautifulSoup(block_html, 'html.parser')
            element = soup_html.find(tag_name)
            if element:
                text_content = element.get_text(separator='\n').strip()
                # 解析对齐（CSS text-align 或 HTML align 属性）
                alignment = extract_alignment(style_attr)
                # 处理块内每一行
                for text_line in text_content.split('\n'):
                    text_line = text_line.strip()
                    if not text_line:
                        continue
                    p = doc.add_paragraph()
                    parse_text_formatting(p, text_line)
                    set_paragraph_format(p)
                    if alignment is not None:
                        p.paragraph_format.alignment = alignment
                    if not has_seen_h2:
                        has_body_before_first_h2 = True
            i += 1
            continue

        # 分割线（必须在 Markdown 表格检测之前，避免 --- 被误判为表格分隔行）
        if line in ['---', '***', '___']:
            if book_mode:
                # 全书合并：每个分隔符 = 章间断点（用 section break 而非 page break，
                # 配合 sectPr footnotePr numRestart=eachSec 实现每章脚注从 1 重置编号）
                doc.add_section(WD_SECTION.NEW_PAGE)
            elif not has_seen_first_hr:
                # 第一个分隔符视为封面与正文的分界，渲染为分页符
                has_seen_first_hr = True
                doc.add_page_break()
                print("✅ 封面分隔符 → 分页符")
            else:
                add_horizontal_line(doc)
            i += 1
            continue

        # Markdown 表格
        if is_table_row(line):
            table_lines = []
            while i < len(lines) and is_table_row(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            if len(table_lines) >= 2:
                create_word_table(doc, table_lines, md_file_path=md_file_path)
                if not has_seen_h2:
                    has_body_before_first_h2 = True
                print(f"✅ 处理Markdown表格: {len(table_lines)} 行")
            continue
        
        # 任务列表
        if line.startswith('- [ ]') or line.startswith('- [x]') or line.startswith('- [X]'):
            add_task_list(doc, line)
            if not has_seen_h2:
                has_body_before_first_h2 = True
            i += 1
            continue
        
        # 无序列表
        if line.startswith(('- ', '* ', '+ ')):
            add_bullet_list(doc, line)
            if not has_seen_h2:
                has_body_before_first_h2 = True
            i += 1
            continue
        
        # 有序列表
        if re.match(r'^\d+\.\s', line):
            add_numbered_list(doc, line)
            if not has_seen_h2:
                has_body_before_first_h2 = True
            i += 1
            continue
        
        # 引用块
        if line.startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].startswith('>'):
                quote_lines.append(lines[i][1:].strip())
                i += 1
            if quote_lines:
                add_quote(doc, '\n'.join(quote_lines))
                if not has_seen_h2:
                    has_body_before_first_h2 = True
            continue

        # Markdown 图片（含外部URL）
        img_match = re.match(r'^!\[([^\]]*)\]\((.+)\)$', line)
        if img_match:
            alt_text = convert_quotes_to_chinese(img_match.group(1))
            img_raw = img_match.group(2)

            # 解析图片路径：支持 URL、绝对路径、相对 md 文件的相对路径；处理 URL 编码（含 %20、空格、中文）
            img_url = img_raw.strip()
            if not img_url.startswith(('http://', 'https://')):
                # 去掉 markdown 链接中可能附带的 title（"path" title）或锚点
                img_url = img_url.split()[0] if ' ' in img_url else img_url
                img_url = urllib.parse.unquote(img_url)
                # 相对路径：基于 md 文件所在目录解析
                if not os.path.isabs(img_url):
                    md_dir = os.path.dirname(os.path.abspath(md_file_path))
                    candidate = os.path.normpath(os.path.join(md_dir, img_url))
                else:
                    candidate = img_url
            else:
                candidate = img_url

            image = None
            if img_url.startswith(('http://', 'https://')):
                print(f"🖼️  下载外部图片: {alt_text[:40]}...")
                image = download_external_image(img_url)
            elif os.path.exists(candidate):
                try:
                    image = Image.open(candidate)
                    image.load()
                except Exception as e:
                    print(f"⚠️  本地图片加载失败: {candidate} ({e})")
            else:
                print(f"⚠️  本地图片不存在: {candidate}")

            if image:
                insert_image_to_word(doc, image)
                print(f"✅ 插入图片: {alt_text[:50]}")
            else:
                # 降级处理：插入文字占位符
                p = doc.add_paragraph()
                parse_text_formatting(p, f"[图片: {alt_text}]")
                set_paragraph_format(p)
                print(f"⚠️  图片未能加载，已插入文字占位符: {alt_text[:40]}")

            if not has_seen_h2:
                has_body_before_first_h2 = True
            i += 1
            continue
        
        # 标题
        if line.startswith('# '):
            title = convert_quotes_to_chinese(line[2:].strip())
            p = doc.add_paragraph()
            parse_text_formatting(p, title, title_level=1)
            set_paragraph_format(p, title_level=1)
        elif line.startswith('## '):
            title = convert_quotes_to_chinese(line[3:].strip())
            p = doc.add_paragraph()
            parse_text_formatting(p, title, title_level=2)
            set_paragraph_format(p, title_level=2)
            has_seen_h2 = True
        elif line.startswith('### '):
            title = convert_quotes_to_chinese(line[4:].strip())
            p = doc.add_paragraph()
            parse_text_formatting(p, title, title_level=3)
            set_paragraph_format(p, title_level=3)
        elif line.startswith('#### '):
            title = convert_quotes_to_chinese(line[5:].strip())
            p = doc.add_paragraph()
            parse_text_formatting(p, title, title_level=4)
            set_paragraph_format(p, title_level=4)
        else:
            if line:
                p = doc.add_paragraph()
                parse_text_with_footnotes(p, line)
                # 图注（**图 X-X：...** / 图 X-X：...）居中、无首行缩进、小一号字
                if re.match(r'^\*{0,2}图\s*\d+[-－]?\d*\s*[:：]', line):
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    pf = p.paragraph_format
                    pf.first_line_indent = Pt(0)
                    pf.left_indent = Pt(0)
                    pf.space_before = Pt(3)
                    pf.space_after = Pt(8)
                    pf.line_spacing = 1.2
                    for r in p.runs:
                        r.font.size = Pt(10)
                else:
                    set_paragraph_format(p)
                if not has_seen_h2:
                    has_body_before_first_h2 = True
        
        i += 1
    
    # endnote 模式：文档末追加“注释”小节
    fn_manager.append_endnotes_section(doc)

    # 添加页码（仅在没有模板时）
    if not use_template_headers:
        add_page_number(doc)

    # footnote + book 模式：每章脚注从 1 重置编号（per-section numRestart=eachSec）
    if book_mode and notes_mode == 'footnote' and fn_manager.refs:
        set_footnote_restart_per_section(doc)
        print('🔖 已设置每章脚注从 1 重置编号（footnotePr numRestart=eachSec）')

    # 保存文档
    doc.save(output_path)

    # footnote 模式：post-process 注入 footnotes.xml part
    fn_manager.finalize_footnotes_part(output_path)

    # 清理全局 manager
    _active_fn_manager = None

    print(f"✅ Word文档已生成: {output_path}")


# ============================================================================
# CLI 入口
# ============================================================================

def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description='Markdown到Word文档转换工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  %(prog)s input.md
  %(prog)s input.md --preset=academic
  %(prog)s input.md --config=my-config.yaml
  %(prog)s input.md output.docx
  %(prog)s --list-presets
        """
    )
    
    parser.add_argument('input', nargs='?', help='输入的 Markdown 文件')
    parser.add_argument('output', nargs='?', help='输出的 Word 文件')
    parser.add_argument('--preset', '-p', help='使用预设配置', default='legal')
    parser.add_argument('--config', '-c', help='使用自定义配置文件 (YAML格式)')
    parser.add_argument('--list-presets', action='store_true', help='列出所有可用的预设配置')
    parser.add_argument('--template', '-t', help='Word模板文件路径')
    parser.add_argument('--auto-template', action='store_true',
                        help='自动从 assets/templates/ 加载第一个 .docx 模板（默认关闭，避免律所 logo 等视觉元素出现在用户没显式要求的 docx 里）')
    parser.add_argument('--landscape', action='store_true', help='使用横向页面（Landscape）')
    parser.add_argument('--notes', choices=['footnote', 'endnote'], default='footnote',
                        help='脚注模式：footnote=页面脚注(默认)；endnote=文档末尾注(上标编号+末尾列表)')
    parser.add_argument('--book', nargs='+', metavar='MD',
                        help='全书合并导出：多章 md → 单 docx（目录 + 章间分页 + 页眉书名）。如 --book ch01.md ch02.md ...')
    parser.add_argument('-o', '--out', dest='out_file',
                        help='输出路径（与 --book 配合；单文件模式用位置参数 output）')
    
    args = parser.parse_args()
    
    if args.list_presets:
        print("可用的预设配置:")
        presets = list_presets()
        if presets:
            for preset in presets:
                cfg = get_preset(preset)
                if cfg:
                    print(f"  - {preset}: {cfg.description}")
        else:
            print("  没有可用的预设配置")
        return
    
    config = None
    if args.config:
        config = load_config(args.config)
        if config is None:
            print(f"❌ 无法加载配置文件: {args.config}")
            return
        print(f"📋 使用配置文件: {args.config}")
    elif args.preset:
        config = get_preset(args.preset)
        if config is None:
            print(f"❌ 预设不存在: {args.preset}")
            print(f"可用预设: {', '.join(list_presets())}")
            return
        print(f"📋 使用预设: {args.preset}")
    
    if config is None:
        config = get_default_preset()
    
    set_config(config)

    if args.book:
        output_file = args.out_file or 'book.docx'
        create_book(args.book, output_file, config, args.notes)
        return

    if not args.input:
        auto_mode(config)
        return
    
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_file = args.input
    if not os.path.isabs(md_file):
        alt = os.path.join(script_dir, md_file)
        if os.path.exists(alt):
            md_file = alt
    
    if not os.path.exists(md_file):
        print(f"❌ 错误: 找不到文件 {md_file}")
        return
    
    output_file = args.output if args.output else generate_output_filename(md_file)
    if args.template:
        template_file = args.template
    else:
        template_file = find_template_file(auto=args.auto_template)

    if args.landscape:
        if 'page' not in config._config:
            config._config['page'] = {}
        config._config['page']['orientation'] = 'landscape'

    try:
        create_word_document(md_file, output_file, template_file, config, args.notes)
        print_success_info(output_file, config)
    except Exception as e:
        print(f"❌ 错误: {e}")
        import traceback
        traceback.print_exc()


def auto_mode(config: Config):
    """自动模式：处理当前目录下的所有.md文件"""
    md_files = find_md_files()
    
    if not md_files:
        print("❌ 当前目录下没有找到.md文件")
        print("\n💡 使用方法:")
        print("1. 将此脚本放在包含.md文件的文件夹中")
        print("2. 或者运行: python md2word.py 输入文件.md")
        print("3. 或者运行: python md2word.py 输入文件.md --preset=academic")
        print("\n📋 可用预设:")
        presets = list_presets()
        if presets:
            for preset in presets:
                cfg = get_preset(preset)
                if cfg:
                    print(f"  - {preset}: {cfg.description}")
        return
    
    print(f"🔍 找到 {len(md_files)} 个Markdown文件:")
    for i, md_file in enumerate(md_files, 1):
        print(f"  {i}. {md_file}")
    
    print("\n开始转换...")
    
    template_file = find_template_file()
    success_count = 0
    
    for md_file in md_files:
        output_file = generate_output_filename(md_file)
        try:
            create_word_document(md_file, output_file, template_file, config)
            success_count += 1
        except Exception as e:
            print(f"❌ 处理 {md_file} 时出错: {e}")
    
    print(f"\n✅ 转换完成！成功处理 {success_count}/{len(md_files)} 个文件")
    print_success_info(None, config)


def print_success_info(filename=None, config: Config = None):
    """打印成功信息"""
    if config is None:
        config = get_config()
    
    print("\n📋 自动应用的格式:")
    
    page_config = config.get('page', {})
    print(f"📄 页面大小: {page_config.get('width', 21.0)}cm × {page_config.get('height', 29.7)}cm")
    print(f"📐 页边距: 上下{page_config.get('margin_top', 2.54)}cm，左右{page_config.get('margin_left', 3.18)}cm")
    
    font_config = config.get('fonts.default', {})
    print(f"📝 字体: {font_config.get('name', '仿宋_GB2312')}")
    print(f"📏 字号: {font_config.get('size', 12)}pt")
    
    paragraph_config = config.get('paragraph', {})
    print(f"📐 行距: {paragraph_config.get('line_spacing', 1.5)}倍")
    
    title1_config = config.get('titles.level1', {})
    print(f"🎯 一级标题: {title1_config.get('size', 15)}pt，{'加粗' if title1_config.get('bold') else '常规'}")
    
    page_number_config = config.get('page_number', {})
    if page_number_config.get('enabled', True):
        print(f"📄 页码设置: {page_number_config.get('format', '1/x')}格式")
    
    quotes_config = config.get('quotes', {})
    if quotes_config.get('convert_to_chinese', True):
        print("💬 引号转换: 英文引号自动转为中文引号")
    
    print("📊 表格支持: Markdown表格自动转换")
    print("📈 图表支持: Mermaid图表本地渲染")
    print("✨ 格式支持: **加粗**、*斜体*、<u>下划线</u>、~~删除线~~")
    print("\n🎯 完全无需手动调整！直接可用！")
    
    if filename:
        print(f"\n📁 输出文件: {filename}")


if __name__ == "__main__":
    main()

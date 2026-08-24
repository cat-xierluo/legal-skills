#!/usr/bin/env python3
"""md2word 已知出版逃逸的回归测试。"""

from pathlib import Path
from tempfile import TemporaryDirectory
import sys
import unittest
import zipfile

from docx import Document
from docx.oxml.ns import qn


HERE = Path(__file__).resolve().parent
if str(HERE) not in sys.path:
    sys.path.insert(0, str(HERE))

from formatter import convert_quotes_to_chinese  # noqa: E402
from footnote_handler import (  # noqa: E402
    FootnoteManager,
    _footnote_text_to_runs_xml,
    _inject_footnotes_into_docx,
)

import md2word  # noqa: E402


class Md2WordRegressionTest(unittest.TestCase):
    def test_cjk_ascii_quotes_convert_but_english_apostrophes_survive(self):
        converted = convert_quotes_to_chinese("标注'需律师现场确认'，don't、O'Brien 与 API's 保留。")
        self.assertIn("‘需律师现场确认’", converted)
        self.assertIn("don't", converted)
        self.assertIn("O'Brien", converted)
        self.assertIn("API's", converted)

    def test_footnote_inline_markers_become_word_properties(self):
        xml = _footnote_text_to_runs_xml("*模型概览* 与 **重点**，命令 `book-gate verify`")
        self.assertNotIn("*模型概览*", xml)
        self.assertNotIn("**重点**", xml)
        self.assertNotIn("`book-gate verify`", xml)
        self.assertIn("<w:i/>", xml)
        self.assertIn("<w:b/>", xml)
        self.assertIn('w:ascii="Consolas"', xml)

    def test_injected_footnotes_xml_has_no_literal_markdown(self):
        with TemporaryDirectory() as temp:
            docx_path = Path(temp) / "footnotes.docx"
            Document().save(docx_path)
            _inject_footnotes_into_docx(
                str(docx_path),
                [(1, "*模型概览*"), (2, "**需律师确认**"), (3, "`Skill`")],
            )
            with zipfile.ZipFile(docx_path) as archive:
                xml = archive.read("word/footnotes.xml").decode("utf-8")
            self.assertNotIn("*模型概览*", xml)
            self.assertNotIn("**需律师确认**", xml)
            self.assertNotIn("`Skill`", xml)
            self.assertIn("<w:i/>", xml)
            self.assertIn("<w:b/>", xml)

    def test_endnotes_path_also_removes_markdown_markers(self):
        document = Document()
        manager = FootnoteManager(mode="endnote")
        manager.refs = [(1, "*模型概览* 与 **重点**，命令 `Skill`")]
        manager.append_endnotes_section(document)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("模型概览 与 重点，命令 Skill", text)
        self.assertNotIn("*", text)
        self.assertNotIn("`", text)

    def test_external_image_download_function_exists(self):
        # 外链图片下载保持默认启用（原行为），download_external_image 可被直接调用
        self.assertTrue(callable(md2word.download_external_image))
        self.assertFalse(hasattr(md2word, "ALLOW_REMOTE_IMAGES"), "外链图片下载开关已移除，保持默认下载")

    def test_book_mode_keeps_in_chapter_hr_and_only_breaks_between_chapters(self):
        with TemporaryDirectory() as temp:
            temp_dir = Path(temp)
            chapter_one = temp_dir / "ch01.md"
            chapter_two = temp_dir / "ch02.md"
            output = temp_dir / "book.docx"
            chapter_one.write_text(
                "# 第一章\n\n正文一[^note]\n\n---\n\n章内分隔线后的正文。\n\n[^note]: 第一章脚注\n",
                encoding="utf-8",
            )
            chapter_two.write_text(
                "# 第二章\n\n正文二[^note]\n\n[^note]: 第二章脚注\n",
                encoding="utf-8",
            )
            config = md2word.get_preset("book-publish")
            md2word.set_config(config)

            md2word.create_book([str(chapter_one), str(chapter_two)], str(output), config)

            document = Document(output)
            horizontal_rule = "─" * config.get("horizontal_rule.repeat_count", 55)
            self.assertEqual(len(document.sections), 2, "两章合并应恰好产生两个 section")
            self.assertEqual(
                sum(paragraph.text == horizontal_rule for paragraph in document.paragraphs),
                1,
                "第一章内的 Markdown 水平线应保留",
            )
            for section in document.sections:
                footnote_properties = section._sectPr.find(qn("w:footnotePr"))
                self.assertIsNotNone(footnote_properties)
                restart = footnote_properties.find(qn("w:numRestart"))
                self.assertEqual(restart.get(qn("w:val")), "eachSec")
            self.assertIn("法律 AI Skill 实战", document.sections[0].header.paragraphs[0].text)
            self.assertTrue(document.sections[1].header.is_linked_to_previous)
            with zipfile.ZipFile(output) as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")
                settings_xml = archive.read("word/settings.xml").decode("utf-8")
            self.assertIn('TOC \\o "1-3" \\h \\z \\u', document_xml)
            self.assertIn("<w:updateFields", settings_xml)

    def test_single_document_first_hr_is_rendered_without_page_break(self):
        with TemporaryDirectory() as temp:
            temp_dir = Path(temp)
            markdown = temp_dir / "single.md"
            output = temp_dir / "single.docx"
            markdown.write_text(
                "# 单章\n\n首段。\n\n---\n\n第二段。\n\n***\n\n第三段。\n\n___\n\n末段。\n",
                encoding="utf-8",
            )
            config = md2word.get_preset("book-publish")
            md2word.set_config(config)

            md2word.create_word_document(str(markdown), str(output), config=config)

            document = Document(output)
            horizontal_rule = "─" * config.get("horizontal_rule.repeat_count", 55)
            self.assertEqual(len(document.sections), 1)
            self.assertEqual(
                sum(paragraph.text == horizontal_rule for paragraph in document.paragraphs),
                3,
                "---、***、___ 都应按 Markdown 语义渲染为水平线",
            )
            with zipfile.ZipFile(output) as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")
            self.assertNotIn('<w:br w:type="page"', document_xml)


if __name__ == "__main__":
    unittest.main(verbosity=2)

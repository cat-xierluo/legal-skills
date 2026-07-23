#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
脚注/尾注处理模块（md2word）

支持两种模式（CLI --notes=footnote|endnote，默认 footnote）：
- footnote：Word 原生页面脚注（正文 footnoteReference + 注入 word/footnotes.xml）
- endnote：文档末“注释”小节 + 正文上标编号（伪 endnote）
  （Word 原生 endnote 只能放文档末、不能“每章末”，故每章尾注用上标编号 + 末尾列表实现）

markdown 语法：
- 引用：[^id]（行内，可多次引用同一 id）
- 定义：[^id]: 脚注文本（单独行，预处理时提取并从正文移除）
"""

import re
import shutil
import zipfile

from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Pt

FOOTNOTES_CT = 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml'
FOOTNOTES_REL = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'

NOTE_DEF_RE = re.compile(r'^\[\^([^\]]+)\]:\s*(.*)$')   # [^id]: text
NOTE_REF_RE = re.compile(r'\[\^([^\]]+)\]')              # [^id] 引用


def extract_footnote_defs(lines):
    """从 md 行列表提取脚注定义。返回 (defs_dict, cleaned_lines)。

    defs_dict: {id: text}；cleaned_lines: 移除定义行后的行列表（其他行原样保留）。
    """
    defs = {}
    cleaned = []
    for line in lines:
        m = NOTE_DEF_RE.match(line.strip())
        if m:
            defs[m.group(1)] = m.group(2).strip()
        else:
            cleaned.append(line)
    return defs, cleaned


class FootnoteManager:
    """管理一次转换的脚注/尾注状态。"""

    def __init__(self, mode='footnote'):
        assert mode in ('footnote', 'endnote')
        self.mode = mode
        self.defs = {}        # {id: text}
        self.refs = []        # [(seq, text), ...] 按出现顺序（去重后每 note_id 一条）
        self._seq = 0         # 编号 / footnote w:id 计数（从 1 起）
        self._id_map = {}     # {note_id: seq} 同一 note_id 复用（标准多次引用同一脚注）

    def set_defs(self, defs):
        self.defs = defs

    def add_reference(self, paragraph, note_id):
        """在段落插入脚注引用。footnote→footnoteReference；endnote→上标编号。

        同一 note_id 多次引用时复用同一 seq：footnotes.xml 只登记一条，正文
        多个 footnoteReference 指向同一 w:id——Word 自动渲染为同号、脚注块
        只出现一次（标准 markdown 多次引用同一脚注的语义）。
        """
        text = self.defs.get(note_id, '')
        if not text:
            return  # 无定义的悬空引用跳过
        if note_id in self._id_map:
            seq = self._id_map[note_id]   # 复用：不重复建 footnotes.xml 条目
        else:
            self._seq += 1
            seq = self._seq
            self._id_map[note_id] = seq
            self.refs.append((seq, text))  # 仅首次引用登记
        if self.mode == 'footnote':
            # 正文 run：footnoteReference（w:id 与 footnotes.xml 中 w:footnote 对应）
            run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            vert = OxmlElement('w:vertAlign')
            vert.set(qn('w:val'), 'superscript')
            rPr.append(vert)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '18')
            rPr.append(sz)
            run.append(rPr)
            fn_ref = OxmlElement('w:footnoteReference')
            fn_ref.set(qn('w:id'), str(seq))
            run.append(fn_ref)
            paragraph._p.append(run)
        else:  # endnote：上标编号
            r = paragraph.add_run(str(seq))
            r.font.superscript = True
            r.font.size = Pt(9)

    def append_endnotes_section(self, doc):
        """endnote 模式：文档末追加“注释”小节 + 编号列表。footnote 模式无操作。"""
        if self.mode != 'endnote' or not self.refs:
            return
        doc.add_paragraph()  # 空行分隔
        title = doc.add_paragraph()
        tr = title.add_run('注释')
        tr.bold = True
        tr.font.size = Pt(12)
        for seq, text in self.refs:
            p = doc.add_paragraph()
            idx = p.add_run('[%d] ' % seq)
            idx.font.size = Pt(9)
            idx.bold = True
            _append_note_runs(p, text)

    def finalize_footnotes_part(self, output_path):
        """footnote 模式：doc.save 后 post-process，注入 footnotes.xml + rels + Content_Types。"""
        if self.mode != 'footnote' or not self.refs:
            return
        _inject_footnotes_into_docx(output_path, self.refs)


def set_footnote_restart_per_section(doc):
    """给 doc 每个 section 的 sectPr 注入 footnotePr numRestart=eachSec，
    实现「每章脚注从 1 重置编号」。

    前提：每章是独立 section（book 模式下 --- 走 doc.add_section 而非 add_page_break）。
    OOXML CT_SectPr 序列里 footnotePr 是第一个子元素，故 insert(0, ...)；
    已存在 footnotePr 时仅覆盖/补 numRestart，不重复注入。
    """
    for section in doc.sections:
        sectPr = section._sectPr
        fnPr = sectPr.find(qn('w:footnotePr'))
        if fnPr is None:
            fnPr = OxmlElement('w:footnotePr')
            sectPr.insert(0, fnPr)
        restart = fnPr.find(qn('w:numRestart'))
        if restart is None:
            restart = OxmlElement('w:numRestart')
            fnPr.append(restart)
        restart.set(qn('w:val'), 'eachSec')


def _xml_escape(t):
    return t.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')


# 脚注 markdown inline 强调 token（**bold** / *italic* / `code`）；
# 不处理 _italic_（避免下划线变量名误判），不处理嵌套，不处理 [link](url)（留 follow-up）。
_FN_INLINE_TOKEN_RE = re.compile(r'(\*\*.+?\*\*|\*.+?\*|`.+?`)')


def _parse_footnote_inline(text):
    """返回 (text, bold, italic, code) 片段，供 footnote/endnote 两条路径共用。"""
    segments = []
    for part in _FN_INLINE_TOKEN_RE.split(text):
        if not part:
            continue
        bold = italic = code = False
        segment = part
        if len(part) >= 4 and part.startswith('**') and part.endswith('**'):
            bold = True
            segment = part[2:-2]
        elif len(part) >= 2 and part.startswith('*') and part.endswith('*'):
            italic = True
            segment = part[1:-1]
        elif len(part) >= 2 and part.startswith('`') and part.endswith('`'):
            code = True
            segment = part[1:-1]
        segments.append((segment, bold, italic, code))
    return segments


def _append_note_runs(paragraph, text):
    """在 python-docx 段落中写入已解析的 endnote runs，不显示 Markdown 标记。"""
    for segment, bold, italic, code in _parse_footnote_inline(text):
        run = paragraph.add_run(segment)
        run.font.size = Pt(9)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = 'Consolas'
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '等线')


def _footnote_text_to_runs_xml(text):
    """把脚注 text 的 markdown inline 强调（**bold**/*italic*/`code`）转为 Word runs XML 片段。

    每个 segment 一个 <w:r>，base rPr 含 sz=18（脚注小字）；bold 加 <w:b/>，
    italic 加 <w:i/>，code 加 Consolas 等宽字体。第一个 run 文本前缀一个空格
    （对齐原 ' %s' 前导空格，与 footnoteRef 拉开间距）。返回 runs XML（不含 <w:p>）。

    bug fix：原实现把 text 直接 _xml_escape 塞进单个 <w:t>，*需律师现场确认* 的
    星号原样进 XML，Word 显示字面星号。本函数解析强调标记转 run，既不显示字面
    星号、又保留斜体/粗体/代码格式。
    """
    runs = []
    first = True
    for seg, bold, italic, code in _parse_footnote_inline(text):
        rpr_parts = ['<w:sz w:val="18"/>']
        if bold:
            rpr_parts.append('<w:b/>')
        if italic:
            rpr_parts.append('<w:i/>')
        if code:
            rpr_parts.append('<w:rFonts w:ascii="Consolas" w:hAnsi="Consolas"/>')
        rpr = '<w:rPr>' + ''.join(rpr_parts) + '</w:rPr>'
        prefix = ' ' if first else ''
        first = False
        runs.append('<w:r>%s<w:t xml:space="preserve">%s%s</w:t></w:r>'
                    % (rpr, prefix, _xml_escape(seg)))
    if not runs:
        runs.append('<w:r><w:rPr><w:sz w:val="18"/></w:rPr>'
                    '<w:t xml:space="preserve"> </w:t></w:r>')
    return ''.join(runs)


def _inject_footnotes_into_docx(docx_path, refs):
    """post-process：把脚注 part 注入已保存的 docx（纯 zip 操作，不依赖 python-docx part API）。

    refs: [(seq, text), ...]，seq 即 footnote w:id。
    footnotes.xml 采用自包含内联格式（vertAlign/sz 内联），不依赖 styles.xml 的
    FootnoteText / FootnoteReference 样式定义，Word 可直接正确渲染。
    """
    tmp_path = docx_path + '.fn_tmp'
    with zipfile.ZipFile(docx_path, 'r') as zin:
        names = zin.namelist()
        data = {n: zin.read(n) for n in names}

    ct = data.get('[Content_Types].xml', b'').decode('utf-8')
    rels_name = 'word/_rels/document.xml.rels'
    rels = data.get(rels_name, b'').decode('utf-8') if rels_name in data else None

    # 1) 构建 footnotes.xml
    fn_items = [
        '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>',
        '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>',
    ]
    for seq, text in refs:
        fn_items.append(
            '<w:footnote w:id="%d"><w:p>'
            '<w:r><w:rPr><w:vertAlign w:val="superscript"/><w:sz w:val="18"/></w:rPr><w:footnoteRef/></w:r>'
            '%s'
            '</w:p></w:footnote>' % (seq, _footnote_text_to_runs_xml(text))
        )
    footnotes_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        + ''.join(fn_items) + '</w:footnotes>'
    )

    # 2) [Content_Types].xml 加 override（若缺）
    if 'footnotes' not in ct:
        override = '<Override PartName="/word/footnotes.xml" ContentType="%s"/>' % FOOTNOTES_CT
        ct = ct.replace('</Types>', override + '</Types>')

    # 3) document.xml.rels 加关系（若缺）
    max_id = 0
    if rels:
        for m in re.finditer(r'Id="rId(\d+)"', rels):
            max_id = max(max_id, int(m.group(1)))
    if rels is None:
        rels = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
                '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                '<Relationship Id="rId%d" Type="%s" Target="footnotes.xml"/></Relationships>'
                % (max_id + 1, FOOTNOTES_REL))
    elif FOOTNOTES_REL not in rels:
        rel_entry = '<Relationship Id="rId%d" Type="%s" Target="footnotes.xml"/>' % (max_id + 1, FOOTNOTES_REL)
        rels = rels.replace('</Relationships>', rel_entry + '</Relationships>')

    data['[Content_Types].xml'] = ct.encode('utf-8')
    data[rels_name] = rels.encode('utf-8')
    data['word/footnotes.xml'] = footnotes_xml.encode('utf-8')

    # 4) 重写 docx（保持原条目顺序，追加 footnotes.xml）
    with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for n in data:
            zout.writestr(n, data[n])
    shutil.move(tmp_path, docx_path)
    print('✅ 已注入 %d 个页面脚注（footnotes.xml）' % len(refs))

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
表格处理模块
处理 Markdown 表格和 HTML 表格的解析与转换
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.shared import OxmlElement
from bs4 import BeautifulSoup, NavigableString

# 导入配置模块
from config import Config, get_config


def set_cell_background_color(cell, color_hex):
    """设置单元格背景色"""
    if not color_hex:
        return
    color = color_hex.lstrip('#')
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _apply_rounded_corners(table, border_color, border_width):
    """为表格四角单元格应用圆角边框效果。

    Word 原生不支持 CSS 圆角，通过为四角单元格单独设置 tcBorders
    并使用稍浅/稍细的外边框来模拟圆角视觉。外部边框用细线、内部网格
    保持原边框粗细，形成层次感。

    Args:
        table: docx Table 对象
        border_color: 十六进制颜色（不含#）
        border_width: 边框粗细（1/8 pt 单位）
    """
    if not table.rows or not table.columns:
        return
    num_rows = len(table.rows)
    num_cols = len(table.columns)
    # 四角单元格坐标
    corners = [
        (0, 0),                       # 左上
        (0, num_cols - 1),           # 右上
        (num_rows - 1, 0),           # 左下
        (num_rows - 1, num_cols - 1)  # 右下
    ]
    # 圆角视觉：外部边框略细，使用稍浅的颜色
    outer_width = max(1, int(border_width * 0.6))
    # 使用稍浅的边框色模拟圆角的柔和过渡
    light_color = _lighten_color(border_color, 0.35)

    for ri, ci in corners:
        cell = table.cell(ri, ci)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        # 移除已有的 tcBorders
        for old in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(old)
        # 构建 tcBorders：外部边框用浅色细线，内部保持原样
        borders_xml = (
            f'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        )
        # 左上角：top + left 用浅色
        if ri == 0 and ci == 0:
            borders_xml += f'<w:top w:val="single" w:sz="{outer_width}" w:space="0" w:color="{light_color}"/>'
            borders_xml += f'<w:left w:val="single" w:sz="{outer_width}" w:space="0" w:color="{light_color}"/>'
        # 右上角：top + right 用浅色
        elif ri == 0 and ci == num_cols - 1:
            borders_xml += f'<w:top w:val="single" w:sz="{outer_width}" w:space="0" w:color="{light_color}"/>'
            borders_xml += f'<w:right w:val="single" w:sz="{outer_width}" w:space="0" w:color="{light_color}"/>'
        # 左下角：bottom + left 用浅色
        elif ri == num_rows - 1 and ci == 0:
            borders_xml += f'<w:bottom w:val="single" w:sz="{outer_width}" w:space="0" w:color="{light_color}"/>'
            borders_xml += f'<w:left w:val="single" w:sz="{outer_width}" w:space="0" w:color="{light_color}"/>'
        # 右下角：bottom + right 用浅色
        elif ri == num_rows - 1 and ci == num_cols - 1:
            borders_xml += f'<w:bottom w:val="single" w:sz="{outer_width}" w:space="0" w:color="{light_color}"/>'
            borders_xml += f'<w:right w:val="single" w:sz="{outer_width}" w:space="0" w:color="{light_color}"/>'
        borders_xml += f'</w:tcBorders>'
        try:
            tcPr.append(parse_xml(borders_xml))
        except Exception:
            pass


def _lighten_color(hex_color, factor=0.35):
    """将十六进制颜色变亮（向白色混合）"""
    hex_color = hex_color.lstrip('#')
    r = int(hex_color[0:2], 16)
    g = int(hex_color[2:4], 16)
    b = int(hex_color[4:6], 16)
    r = min(255, int(r + (255 - r) * factor))
    g = min(255, int(g + (255 - g) * factor))
    b = min(255, int(b + (255 - b) * factor))
    return f'{r:02X}{g:02X}{b:02X}'


def is_separator_line(line):
    """判断是否是表格分隔行。分隔行必须包含'-'，且只能包含'|', '-', ':', ' '等符号。"""
    line = line.strip()
    if not line or '-' not in line:
        return False
    return all(c in '|-: 	' for c in line)


def is_table_row(line):
    """判断是否是表格行"""
    if not line or not line.strip():
        return False

    line = line.strip()

    # 检查是否是分隔行
    if is_separator_line(line):
        return True

    # 检查是否是数据行（包含 |）
    # 这里的逻辑保持宽松，依赖于主循环中对其他块级元素的优先判断
    if '|' in line:
        return True

    return False


def create_word_table(doc, table_lines, md_file_path=None):
    """从Markdown表格行创建Word表格"""

    if len(table_lines) < 2:
        return

    # 解析表格数据
    rows_data = []
    header_row = None

    for i, line in enumerate(table_lines):
        # 跳过分隔行（包含横线的行）
        if is_separator_line(line):
            continue

        # 解析单元格
        cells = parse_table_row(line)
        if cells:
            if header_row is None:
                header_row = cells
            else:
                rows_data.append(cells)

    if not header_row:
        return

    # 确定列数
    max_cols = len(header_row)
    for row in rows_data:
        max_cols = max(max_cols, len(row))

    # 创建Word表格
    total_rows = 1 + len(rows_data)  # 标题行 + 数据行
    table = doc.add_table(rows=total_rows, cols=max_cols)

    # 获取表格配置
    config = get_config()
    table_config = config.get('table', {})
    border_enabled = table_config.get('border_enabled', True)
    border_color = table_config.get('border_color', '#000000')
    border_width = table_config.get('border_width', 4)
    row_height_cm = table_config.get('row_height_cm', 0.8)
    alignment_str = table_config.get('alignment', 'center')
    line_spacing = table_config.get('line_spacing', 1.2)
    cell_margin = table_config.get('cell_margin', {})
    vertical_align_str = table_config.get('vertical_align', 'center')

    # 设置表格对齐方式
    alignment_map = {
        'left': WD_TABLE_ALIGNMENT.LEFT,
        'center': WD_TABLE_ALIGNMENT.CENTER,
        'right': WD_TABLE_ALIGNMENT.RIGHT
    }
    table.alignment = alignment_map.get(alignment_str.lower(), WD_TABLE_ALIGNMENT.CENTER)

    # 设置垂直对齐
    vertical_align_map = {
        'top': WD_ALIGN_VERTICAL.TOP,
        'center': WD_ALIGN_VERTICAL.CENTER,
        'bottom': WD_ALIGN_VERTICAL.BOTTOM
    }
    vertical_align = vertical_align_map.get(vertical_align_str.lower(), WD_ALIGN_VERTICAL.CENTER)

    # 统一设置边框和内边距、行高等
    if border_enabled:
        try:
            tbl = table._tbl
            color = border_color.lstrip('#')
            borders_xml = f'''
            <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:left w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:bottom w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:right w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:insideH w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
                <w:insideV w:val="single" w:sz="{border_width}" w:space="0" w:color="{color}"/>
            </w:tblBorders>
            '''
            tbl.tblPr.append(parse_xml(borders_xml))
        except Exception:
            pass

    try:
        tbl = table._tbl
        top = cell_margin.get('top', 30)
        bottom = cell_margin.get('bottom', 30)
        left = cell_margin.get('left', 60)
        right = cell_margin.get('right', 60)
        cell_margins_xml = f'''
        <w:tblCellMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:w="{top}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tblCellMar>
        '''
        tbl.tblPr.append(parse_xml(cell_margins_xml))
    except Exception:
        pass

    # 行高与段落行距统一
    try:
        for row in table.rows:
            row.height = Cm(row_height_cm)
            for cell in row.cells:
                cell.vertical_alignment = vertical_align
                for paragraph in cell.paragraphs:
                    pf = paragraph.paragraph_format
                    pf.line_spacing = line_spacing
                    pf.space_before = Pt(2)
                    pf.space_after = Pt(2)
    except Exception:
        pass

    # 获取表头背景色配置
    header_bg_color = config.get('table.header', {}).get('background_color')

    # 填充标题行
    header_cells = table.rows[0].cells
    for j, cell_text in enumerate(header_row):
        if j < len(header_cells):
            cell = header_cells[j]
            # 处理表格单元格中的格式
            if contains_markdown_formatting(cell_text.strip()):
                parse_table_cell_formatting(cell, cell_text.strip(), is_header=True, md_file_path=md_file_path)
            else:
                # 导入 convert_quotes_to_chinese 避免循环导入
                from formatter import convert_quotes_to_chinese
                cell.text = convert_quotes_to_chinese(cell_text.strip())
                set_table_cell_format(cell, is_header=True)
            # 应用表头背景色
            if header_bg_color:
                set_cell_background_color(cell, header_bg_color)

    # 获取交替行颜色配置
    row_even_color = config.get('table.row_even', {}).get('background_color')
    row_odd_color = config.get('table.row_odd', {}).get('background_color')

    # 填充数据行
    for i, row_data in enumerate(rows_data):
        if i + 1 < len(table.rows):
            row_cells = table.rows[i + 1].cells
            # 确定当前行颜色（奇偶交替）
            row_bg_color = row_odd_color if i % 2 == 0 else row_even_color
            for j, cell_text in enumerate(row_data):
                if j < len(row_cells):
                    cell = row_cells[j]
                    # 处理表格单元格中的格式
                    if contains_markdown_formatting(cell_text.strip()):
                        parse_table_cell_formatting(cell, cell_text.strip(), is_header=False, md_file_path=md_file_path)
                    else:
                        # 导入 convert_quotes_to_chinese 避免循环导入
                        from formatter import convert_quotes_to_chinese
                        cell.text = convert_quotes_to_chinese(cell_text.strip())
                        set_table_cell_format(cell, is_header=False)
                    # 应用交替行背景色
                    if row_bg_color:
                        set_cell_background_color(cell, row_bg_color)

    # 调整列宽
    adjust_table_column_width(table)

    # 圆角边框效果
    if table_config.get('rounded_corners', False) and border_enabled:
        _apply_rounded_corners(table, border_color.lstrip('#'), border_width)


def parse_table_row(line):
    """解析表格行，提取单元格内容"""
    if not line or not line.strip():
        return []

    line = line.strip()

    # 移除开头和结尾的 |
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]

    # 分割单元格
    cells = [cell.strip() for cell in line.split('|')]

    # 过滤掉空单元格（但保留有意义的空单元格）
    return cells


def contains_markdown_formatting(text):
    """检查文本是否包含Markdown格式标记"""
    format_patterns = [
        r'!\[.*?\]\([^)]+\)',  # 图片语法
        r'\*\*\*.*?\*\*\*',  # 加粗斜体
        r'\*\*.*?\*\*',      # 加粗
        r'\*.*?\*',          # 斜体
        r'___.*?___',        # 加粗斜体
        r'__.*?__',          # 加粗
        r'_.*?_',            # 斜体
        r'<u>.*?</u>',       # 下划线
        r'<strong>.*?</strong>',  # HTML 加粗
        r'<b>.*?</b>',       # HTML 加粗
        r'<em>.*?</em>',     # HTML 斜体
        r'<i>.*?</i>',       # HTML 斜体
        r'<s>.*?</s>',       # HTML 删除线
        r'<del>.*?</del>',   # HTML 删除线
        r'<strike>.*?</strike>',  # HTML 删除线
        r'~~.*?~~',          # 删除线
        r'`.*?`',            # 行内代码
        r'<br\s*/?>',       # 换行标签
        r'\$.*?\$',         # LaTeX数学公式
    ]

    for pattern in format_patterns:
        if re.search(pattern, text):
            return True
    return False


def parse_table_cell_formatting(cell, text, is_header=False, md_file_path=None):
    """解析表格单元格中的格式化文本（含图片）"""
    # 清空单元格
    cell.text = ""

    # 设置段落对齐方式（与 set_table_cell_format 保持一致）
    config = get_config()
    table_config = config.get('table', {})
    line_spacing = table_config.get('line_spacing', 1.2)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = line_spacing

    # 导入 convert_quotes_to_chinese 和 parse_formatted_text 避免循环导入
    from formatter import convert_quotes_to_chinese, parse_formatted_text

    # 转换引号
    text = convert_quotes_to_chinese(text)

    # 先把 markdown 图片语法 ![alt](path) 提取出来，单独处理；剩余文本按格式处理
    img_pattern = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')
    matches = list(img_pattern.finditer(text))

    if matches and md_file_path:
        # 把文本按图片位置切成"图片前文本/图片/图片后文本/..."
        cursor = 0
        for m in matches:
            alt = m.group(1)
            raw = m.group(2).strip().split()[0]  # 去掉可能的 title
            # URL 解码 + 相对路径解析（基于 md_file_path 所在目录）
            from urllib.parse import unquote
            raw = unquote(raw)
            if not raw.startswith(('http://', 'https://')):
                if os.path.isabs(raw):
                    candidate = raw
                else:
                    md_dir = os.path.dirname(os.path.abspath(md_file_path))
                    candidate = os.path.normpath(os.path.join(md_dir, raw))
            else:
                candidate = raw

            # 先写图片前的文本
            pre_text = text[cursor:m.start()]
            if pre_text.strip():
                _render_text_into_cell(cell, pre_text, is_header)
            # 插入图片
            if os.path.exists(candidate):
                try:
                    from PIL import Image as _PILImage
                    with _PILImage.open(candidate) as pil_img:
                        # 表格中图片按 130px 宽对应 Cm 计算
                        from docx.shared import Cm as _Cm
                        run = cell.paragraphs[-1].add_run()
                        run.add_picture(candidate, width=_Cm(5.0))
                    # 在图片下另起一段写 alt 文字（作为说明）
                    if alt.strip():
                        cap_p = cell.add_paragraph()
                        cap_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        cap_run = cap_p.add_run(alt)
                        set_table_run_format(cap_run, {'italic': True}, is_header)
                    print(f"✅ 表格图片插入成功: {os.path.basename(candidate)}")
                except Exception as e:
                    print(f"⚠️  表格图片插入失败: {candidate} ({e})")
                    # 降级：把整段原样写入
                    fallback = cell.add_paragraph()
                    fallback_run = fallback.add_run(m.group(0))
                    set_table_run_format(fallback_run, {}, is_header)
            else:
                print(f"⚠️  表格图片路径不存在: {candidate}")
                # 降级：原样写入
                fallback = cell.add_paragraph()
                fallback_run = fallback.add_run(m.group(0))
                set_table_run_format(fallback_run, {}, is_header)
            cursor = m.end()
        # 写图片后剩余文本
        tail_text = text[cursor:]
        if tail_text.strip():
            _render_text_into_cell(cell, tail_text, is_header)
        return

    # 无图片：走原逻辑
    _render_text_into_cell(cell, text, is_header)


def _render_text_into_cell(cell, text, is_header):
    """将格式化文本（不含图片 markdown 语法）写入表格 cell"""
    from formatter import convert_quotes_to_chinese, parse_formatted_text

    # 转换引号
    text = convert_quotes_to_chinese(text)

    # 支持<br>换行：拆分后逐段处理
    parts_by_br = re.split(r'<br\s*/?>', text, flags=re.IGNORECASE)

    # 解析格式
    format_patterns = [
        (r'\*\*\*(.*?)\*\*\*', {'bold': True, 'italic': True}),
        (r'___(.*?)___', {'bold': True, 'italic': True}),
        (r'\*\*(.*?)\*\*', {'bold': True}),
        (r'__(.*?)__', {'bold': True}),
        (r'(?<!\*)\*([^*\n]+?)\*(?!\*)', {'italic': True}),
        (r'(?<!_)_([^_\n]+?)_(?!_)', {'italic': True}),
        (r'<strong>(.*?)</strong>', {'bold': True}),
        (r'<b>(.*?)</b>', {'bold': True}),
        (r'<em>(.*?)</em>', {'italic': True}),
        (r'<i>(.*?)</i>', {'italic': True}),
        (r'<u>(.*?)</u>', {'underline': True}),
        (r'~~(.*?)~~', {'strikethrough': True}),
        (r'<s>(.*?)</s>', {'strikethrough': True}),
        (r'<del>(.*?)</del>', {'strikethrough': True}),
        (r'<strike>(.*?)</strike>', {'strikethrough': True}),
        (r'`([^`\n]+)`', {'code': True}),
        (r'\$([^$\n]+?)\$', {'math': True}),  # LaTeX数学公式支持
    ]

    for idx, segment in enumerate(parts_by_br):
        if idx > 0:
            cell.paragraphs[0].add_run().add_break()
        text_parts = parse_formatted_text(segment, format_patterns)
        for part_text, formats in text_parts:
            if part_text:  # 只有非空文本才创建run
                run = cell.paragraphs[0].add_run(part_text)
                set_table_run_format(run, formats, is_header)


def set_table_run_format(run, formats, is_header=False):
    """设置表格单元格run格式"""
    config = get_config()

    if is_header:
        header_config = config.get('table.header', {})
        font_name = header_config.get('font', 'Times New Roman')
        font_size = header_config.get('size', 10.5)
        color_hex = header_config.get('color', '#000000')
        bold = header_config.get('bold', True)
    else:
        body_config = config.get('table.body', {})
        font_name = body_config.get('font', '仿宋_GB2312')
        font_size = body_config.get('size', 10.5)
        color_hex = body_config.get('color', '#000000')
        bold = False

    font = run.font
    font.name = 'Times New Roman'  # 默认英文字体
    font.size = Pt(font_size)
    font.color.rgb = hex_to_rgb(color_hex)
    font.bold = bold if is_header else False

    # 设置字体映射：英文和数字用Times New Roman，中文用配置的字体
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run._element.rPr.rFonts.set(qn('w:cs'), 'Times New Roman')
    # 对含 emoji/symbol 的文本使用覆盖更广的字体
    if _contains_symbol_chars(run.text or ''):
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Segoe UI Symbol')
    else:
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')

    # 应用Markdown格式
    if formats.get('bold', False):
        font.bold = True
    if formats.get('italic', False):
        font.italic = True
    if formats.get('underline', False):
        font.underline = True
    if formats.get('strikethrough', False):
        font.strike = True
    if formats.get('code', False):
        # 表格中代码使用Times New Roman，稍小字号
        code_config = config.get('inline_code', {})
        font.name = code_config.get('font', 'Times New Roman')
        font.size = Pt(9)
        font.color.rgb = hex_to_rgb(code_config.get('color', '#333333'))
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        return
    if formats.get('math', False):
        # 表格中数学公式使用Times New Roman，斜体，深蓝色
        math_config = config.get('math', {})
        font.name = math_config.get('font', 'Times New Roman')
        font.size = Pt(math_config.get('size', 10))
        font.italic = math_config.get('italic', True)
        font.color.rgb = hex_to_rgb(math_config.get('color', '#00008B'))
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        return


def set_table_cell_format(cell, is_header=False):
    """设置表格单元格格式"""
    config = get_config()
    table_config = config.get('table', {})
    line_spacing = table_config.get('line_spacing', 1.2)

    if is_header:
        header_config = config.get('table.header', {})
        font_name = header_config.get('font', 'Times New Roman')
        font_size = header_config.get('size', 10.5)
        color_hex = header_config.get('color', '#000000')
        bold = header_config.get('bold', True)
    else:
        body_config = config.get('table.body', {})
        font_name = body_config.get('font', '仿宋_GB2312')
        font_size = body_config.get('size', 10.5)
        color_hex = body_config.get('color', '#000000')
        bold = False

    # 设置段落格式
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = line_spacing

        # 设置文字格式
        for run in paragraph.runs:
            font = run.font
            font.name = font_name
            font.size = Pt(font_size)
            font.color.rgb = hex_to_rgb(color_hex)
            font.bold = bold if is_header else False

            # 设置中文字体
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def adjust_table_column_width(table):
    """基于 P80 百分位按比例分配列宽（Markdown 表格）"""
    try:
        config = get_config()
        col_count = len(table.columns)
        if col_count == 0:
            return

        # 收集每列每个单元格的内容长度（权重字+真实字双口径）
        cell_lengths = {i: [] for i in range(col_count)}        # 权重字(中文×2 ASCII×1)给 P80 瓜分用
        cell_lengths_real = {i: [] for i in range(col_count)}   # 真实字(中文×1 ASCII×0.5)给列类型判定用
        seen = set()
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                if ci >= col_count:
                    continue
                # 不去重: 13行2列无合并, seen 在某些 python-docx 版本下会误判 tc 重复
                text = cell.text.strip()
                w_len = sum(2.0 if ord(c) > 127 else 1.0 for c in text)  # 权重字
                r_len = sum(1 if ord(c) > 127 else 0.5 for c in text)     # 真实字
                cell_lengths[ci].append(max(1.0, w_len))
                cell_lengths_real[ci].append(max(1.0, r_len))

        # 最小必要列宽：按"表头真实字符数"算（不是最长单元格）
        # 原因：表头装得下 → 该列所有短行也装得下（"稳定层"/"业务层"3字与"层级"同列）
        # 长单元格物理上必换行，不作为保底（不然会把整列压出页面）
        # Word 12pt 中文字符约 0.32cm, ASCII 约 0.16cm, 单元格左右内边距 0.84cm
        # 字符数按真实字数算（中文 1 字，英文/数字 0.5 字取整为 1）—— 不再用 P80 的"中文×2"权重（那是相对权重，不是物理字符数）
        cell_margin_cm = 0.84
        min_needed_cm = []
        if table.rows:
            header_cells = table.rows[0].cells
            for ci in range(col_count):
                text = header_cells[ci].text.strip() if ci < len(header_cells) else ''
                # 真实字符数：中文算 1, ASCII 算 0.5, ceil
                real_chars = sum(1 if ord(c) > 127 else 0.5 for c in text)
                real_chars = max(1, int(real_chars + 0.5))
                content_w = real_chars * 0.32  # 中文为主按 0.32/字
                min_needed_cm.append(content_w + cell_margin_cm)
        else:
            min_needed_cm = [1.5] * col_count
        widths_cm = _calc_column_widths(cell_lengths, col_count, config, min_needed_cm=min_needed_cm, cell_lengths_real_per_col=cell_lengths_real)

        # 关键：禁 autofit + fixed layout + 设 gridCol/cell.width，否则 Word 用 autofit 覆盖成等分
        table.autofit = False
        table.allow_autofit = False
        tblPr = table._tbl.tblPr
        for old_layout in tblPr.findall(qn('w:tblLayout')):
            tblPr.remove(old_layout)
        layout = OxmlElement('w:tblLayout')
        layout.set(qn('w:type'), 'fixed')
        tblPr.append(layout)

        # tblGrid 的 gridCol 宽度（cm→twips，1cm≈567twip）
        tblGrid = table._tbl.find(qn('w:tblGrid'))
        if tblGrid is not None:
            gridCols = tblGrid.findall(qn('w:gridCol'))
            for i, w in enumerate(widths_cm):
                if i < len(gridCols):
                    gridCols[i].set(qn('w:w'), str(int(w * 567)))

        # 每个 cell 的 width（双保险，Word 真正按此列宽渲染）
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                if ci < col_count:
                    cell.width = Cm(widths_cm[ci])

        # columns.width（兼容）
        for i, w in enumerate(widths_cm):
            table.columns[i].width = Cm(w)
    except Exception as e:
        print(f"⚠️  列宽调整失败: {e}")


def parse_html_table(html_content):
    """解析HTML表格内容，返回表格数据"""
    try:
        soup = BeautifulSoup(html_content, 'html.parser')
        table = soup.find('table')
        if not table:
            return None

        rows_data = []
        for tr in table.find_all('tr'):
            row_cells = []
            for cell in tr.find_all(['td', 'th']):
                # 获取单元格文本内容，保留基本格式
                cell_text = cell.get_text(strip=True)
                row_cells.append(cell_text)
            if row_cells:  # 只添加非空行
                rows_data.append(row_cells)

        return rows_data
    except Exception as e:
        print(f"⚠️  HTML表格解析失败: {e}")
        return None


def _populate_cell_with_rich_content(cell_element, word_cell, md_file_path=None):
    """将HTML单元格内容填充到Word单元格中，支持图片和富文本"""
    from formatter import convert_quotes_to_chinese
    config = get_config()
    font_config = config.get('fonts.default', {})
    font_name = font_config.get('name', '仿宋_GB2312')
    font_size = font_config.get('size', 10.5)

    # 清空默认段落
    for p in word_cell.paragraphs:
        p.clear()

    first_paragraph = word_cell.paragraphs[0]
    current_paragraph = first_paragraph

    # 检查是否包含图片
    img_tags = cell_element.find_all('img')
    if img_tags:
        for img_tag in img_tags:
            src = img_tag.get('src', '')
            img_width = img_tag.get('width', '130')
            if src and md_file_path:
                # 支持 file:// 协议、绝对路径和相对路径
                if src.startswith('file://'):
                    from urllib.parse import unquote
                    img_path = unquote(src[7:])  # 去掉 file:// 前缀，保留 /path
                elif os.path.isabs(src):
                    img_path = src
                else:
                    md_dir = os.path.dirname(os.path.abspath(md_file_path))
                    img_path = os.path.join(md_dir, src)
                if os.path.exists(img_path):
                    try:
                        # 在当前段落插入图片
                        run = current_paragraph.add_run()
                        run.add_picture(img_path, width=Cm(int(img_width) / 96 * 2.54))
                        # 图片后换新段落
                        current_paragraph = word_cell.add_paragraph()
                        print(f"✅ 表格图片插入成功: {os.path.basename(img_path)}")
                    except Exception as e:
                        print(f"⚠️  表格图片插入失败: {e}")
                else:
                    print(f"⚠️  表格图片路径不存在: {img_path}")
        return

    # 处理文本内容：逐行解析，支持 Markdown 粗体、HTML 内联标签、列表等
    raw_text = cell_element.decode_contents()
    # 先将 <br> 标签转为换行符，确保 HTML 表格中的换行生效
    raw_text = re.sub(r'<br\s*/?>|</br>', '\n', raw_text, flags=re.IGNORECASE)
    lines = raw_text.strip().split('\n')

    # HTML 内联标签格式模式
    html_format_patterns = [
        (r'\*\*\*(.*?)\*\*\*', {'bold': True, 'italic': True}),
        (r'\*\*(.*?)\*\*', {'bold': True}),
        (r'(?<!\*)\*([^*\n]+?)\*(?!\*)', {'italic': True}),
        (r'<strong>(.*?)</strong>', {'bold': True}),
        (r'<b>(.*?)</b>', {'bold': True}),
        (r'<em>(.*?)</em>', {'italic': True}),
        (r'<i>(.*?)</i>', {'italic': True}),
        (r'<u>(.*?)</u>', {'underline': True}),
        (r'<s>(.*?)</s>', {'strikethrough': True}),
        (r'<del>(.*?)</del>', {'strikethrough': True}),
        (r'<strike>(.*?)</strike>', {'strikethrough': True}),
        (r'~~(.*?)~~', {'strikethrough': True}),
    ]

    for line_idx, line in enumerate(lines):
        line_stripped = line.strip()
        if not line_stripped:
            continue

        # 如果不是第一行，添加新段落
        if line_idx > 0:
            current_paragraph = word_cell.add_paragraph()

        # 处理 Markdown 列表
        if line_stripped.startswith('- '):
            line_stripped = line_stripped[2:]
            # 添加列表标记
            run = current_paragraph.add_run('• ')
            run.font.size = Pt(font_size)

        # 使用 parse_formatted_text 处理混合格式
        from formatter import parse_formatted_text
        text_parts = parse_formatted_text(line_stripped, html_format_patterns)
        for part_text, formats in text_parts:
            if part_text:
                run = current_paragraph.add_run(convert_quotes_to_chinese(part_text))
                run.font.size = Pt(font_size)
                run.font.name = font_config.get('ascii', 'Times New Roman')
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                # 对含 emoji/symbol 的文本设置兼容字体
                if _contains_symbol_chars(part_text):
                    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Segoe UI Symbol')
                if formats.get('bold'):
                    run.font.bold = True
                if formats.get('italic'):
                    run.font.italic = True
                if formats.get('underline'):
                    run.font.underline = True
                if formats.get('strikethrough'):
                    run.font.strike = True


def create_word_table_from_html(doc, html_content, md_file_path=None):
    """从HTML表格创建Word表格，支持合并单元格(colspan/rowspan)、边框和列宽优化"""
    try:
        # 去除 HTML 注释
        html_content = re.sub(r'<!--.*?-->', '', html_content, flags=re.DOTALL)

        soup = BeautifulSoup(html_content, 'html.parser')
        html_table = soup.find('table')
        if not html_table:
            print("⚠️  HTML表格数据为空或格式不正确")
            return

        rows = html_table.find_all('tr')
        if not rows:
            print("⚠️  HTML表格无有效行")
            return

        config = get_config()
        table_config = config.get('table', {})

        # ── 第一步：解析网格结构，处理 colspan / rowspan ──
        occupancy = {}  # (row, col) -> cell_info
        row_cells_list = []

        for ri, tr in enumerate(rows):
            cells = tr.find_all(['td', 'th'])
            col_idx = 0
            current_row_cells = []

            for cell in cells:
                # 跳过被上方 rowspan 占据的列
                while (ri, col_idx) in occupancy:
                    col_idx += 1

                colspan = int(cell.get('colspan', 1))
                rowspan = int(cell.get('rowspan', 1))

                cell_info = {
                    'element': cell,
                    'colspan': colspan,
                    'rowspan': rowspan,
                    'is_header': cell.name == 'th',
                    'origin_row': ri,
                    'origin_col': col_idx,
                }
                current_row_cells.append(cell_info)

                # 标记合并区域的所有位置
                for dr in range(rowspan):
                    for dc in range(colspan):
                        occupancy[(ri + dr, col_idx + dc)] = cell_info

                col_idx += colspan

            row_cells_list.append(current_row_cells)

        if not occupancy:
            return

        num_rows = max(r for (r, _) in occupancy) + 1
        num_cols = max(c for (_, c) in occupancy) + 1

        # ── 第二步：创建 Word 表格 ──
        table = doc.add_table(rows=num_rows, cols=num_cols)

        alignment_map = {
            'left': WD_TABLE_ALIGNMENT.LEFT,
            'center': WD_TABLE_ALIGNMENT.CENTER,
            'right': WD_TABLE_ALIGNMENT.RIGHT,
        }
        table.alignment = alignment_map.get(
            table_config.get('alignment', 'center').lower(), WD_TABLE_ALIGNMENT.CENTER
        )

        # ── 第三步：合并单元格 ──
        for row_cells in row_cells_list:
            for ci in row_cells:
                if ci['colspan'] > 1 or ci['rowspan'] > 1:
                    table.cell(ci['origin_row'], ci['origin_col']).merge(
                        table.cell(
                            ci['origin_row'] + ci['rowspan'] - 1,
                            ci['origin_col'] + ci['colspan'] - 1,
                        )
                    )

        # ── 第四步：设置边框和内边距 ──
        _apply_table_borders(table, table_config)
        _apply_table_cell_margins(table, table_config)

        vertical_align_map = {
            'top': WD_ALIGN_VERTICAL.TOP,
            'center': WD_ALIGN_VERTICAL.CENTER,
            'bottom': WD_ALIGN_VERTICAL.BOTTOM,
        }

        # ── 第五步：填充内容并设置格式 ──
        header_bg_color = config.get('table.header', {}).get('background_color')

        for row_cells in row_cells_list:
            for ci in row_cells:
                word_cell = table.cell(ci['origin_row'], ci['origin_col'])

                valign_attr = ci['element'].get('valign', 'center')
                word_cell.vertical_alignment = vertical_align_map.get(
                    valign_attr, WD_ALIGN_VERTICAL.CENTER
                )

                _populate_cell_with_rich_content(ci['element'], word_cell, md_file_path)

                for para in word_cell.paragraphs:
                    pf = para.paragraph_format
                    pf.line_spacing = table_config.get('line_spacing', 1.2)
                    pf.space_before = Pt(2)
                    pf.space_after = Pt(2)
                    if ci['is_header'] or ci['colspan'] > 1:
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                if ci['is_header'] and header_bg_color:
                    set_cell_background_color(word_cell, header_bg_color)

        # ── 第六步：优化列宽 ──
        _optimize_html_table_widths(table, num_cols, row_cells_list, config)

        # 圆角边框效果
        if table_config.get('rounded_corners', False) and table_config.get('border_enabled', True):
            border_color = table_config.get('border_color', '#000000').lstrip('#')
            border_width = table_config.get('border_width', 4)
            _apply_rounded_corners(table, border_color, border_width)

        print(f"✅ 处理HTML表格: {num_rows} 行 x {num_cols} 列")
    except Exception as e:
        print(f"⚠️  HTML表格处理失败: {e}")
        import traceback
        traceback.print_exc()


def _apply_table_borders(table, table_config):
    """为表格添加边框"""
    border_enabled = table_config.get('border_enabled', True)
    if not border_enabled:
        return
    border_color = table_config.get('border_color', '#000000').lstrip('#')
    border_width = table_config.get('border_width', 4)
    try:
        tbl = table._tbl
        borders_xml = (
            f'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f'<w:top w:val="single" w:sz="{border_width}" w:space="0" w:color="{border_color}"/>'
            f'<w:left w:val="single" w:sz="{border_width}" w:space="0" w:color="{border_color}"/>'
            f'<w:bottom w:val="single" w:sz="{border_width}" w:space="0" w:color="{border_color}"/>'
            f'<w:right w:val="single" w:sz="{border_width}" w:space="0" w:color="{border_color}"/>'
            f'<w:insideH w:val="single" w:sz="{border_width}" w:space="0" w:color="{border_color}"/>'
            f'<w:insideV w:val="single" w:sz="{border_width}" w:space="0" w:color="{border_color}"/>'
            f'</w:tblBorders>'
        )
        tbl.tblPr.append(parse_xml(borders_xml))
    except Exception:
        pass


def _apply_table_cell_margins(table, table_config):
    """为表格设置单元格内边距"""
    cell_margin = table_config.get('cell_margin', {})
    try:
        tbl = table._tbl
        top = cell_margin.get('top', 30)
        bottom = cell_margin.get('bottom', 30)
        left = cell_margin.get('left', 60)
        right = cell_margin.get('right', 60)
        cell_margins_xml = (
            f'<w:tblCellMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f'<w:top w:w="{top}" w:type="dxa"/>'
            f'<w:left w:w="{left}" w:type="dxa"/>'
            f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
            f'<w:right w:w="{right}" w:type="dxa"/>'
            f'</w:tblCellMar>'
        )
        tbl.tblPr.append(parse_xml(cell_margins_xml))
    except Exception:
        pass


def _calc_column_widths(cell_lengths_per_col, num_cols, config, min_needed_cm=None, cell_lengths_real_per_col=None):
    """按 P80 百分位给长列宽度 + 短列保底不换行。

    核心策略：
    1. 每列先按"最长单元格实际需要"给一个最低宽(min_needed_cm,核心:少字列不换行)
    2. 长列按 P80 比例瓜分"页面剩余宽度"(sum_min 给定后的余量)
    3. 长列拿到的宽度可能物理上仍不够(必换行) → 用户接受(短列不被压换行是底线)
    """
    page_config = config.get('page', {})
    page_width = page_config.get('width', 21.0)
    margin_left = page_config.get('margin_left', 3.18)
    margin_right = page_config.get('margin_right', 3.18)
    available_cm = page_width - margin_left - margin_right

    if min_needed_cm is None:
        min_needed_cm = [1.0] * num_cols

    # 1) 短列先按 min_needed 给定
    widths_cm = list(min_needed_cm)
    sum_min = sum(widths_cm)

    # 2) 长列按 P80 比例瓜分剩余宽度
    if sum_min < available_cm:
        col_weights = []
        for i in range(num_cols):
            lengths = cell_lengths_per_col.get(i, [])
            if not lengths:
                col_weights.append(0.0)
            else:
                lengths_sorted = sorted(lengths)
                p80_idx = int(len(lengths_sorted) * 0.8)
                col_weights.append(lengths_sorted[min(p80_idx, len(lengths_sorted) - 1)])

        total_weight = sum(col_weights) or 1.0
        extra_cm = available_cm - sum_min

        for i, w in enumerate(col_weights):
            widths_cm[i] = widths_cm[i] + extra_cm * w / total_weight

    # 3) sum_min > available 时(内容总长超页面,物理必换行)
    #    关键:不再整体缩(会把少字列也压扁,这是用户痛点)
    #    而是只缩那些"实际需要 > 最小合理宽"的长列,少字列保底不变
    else:
        # 设一个"合理最大宽"——A4 单栏正文区(14.64cm)的 70% 给单列,让该列内换行可接受
        MAX_REASONABLE = available_cm * 0.7   # ≈ 10.25cm
        # 找出需要压缩的列(min_needed > MAX_REASONABLE)
        excess_total = sum_min - available_cm
        for i, w in enumerate(widths_cm):
            if w > MAX_REASONABLE:
                # 长列:最多压到 MAX_REASONABLE,多余的"还回去"给总宽补足
                can_reduce = w - MAX_REASONABLE
                reduce = min(can_reduce, excess_total)
                widths_cm[i] -= reduce
                excess_total -= reduce
            # 短列(w <= MAX_REASONABLE):保底,不动

    return widths_cm


def _optimize_html_table_widths(table, num_cols, row_cells_list, config):
    """HTML 表格列宽优化"""
    try:
        # 收集每列所有单元格的内容长度
        cell_lengths = {}
        for row_cells in row_cells_list:
            for ci in row_cells:
                text = ci['element'].get_text(strip=True)
                content_len = sum(2.0 if ord(c) > 127 else 1.0 for c in text)
                colspan = ci['colspan']
                origin_c = ci['origin_col']
                if colspan == 1:
                    cell_lengths.setdefault(origin_c, []).append(content_len)
                else:
                    per_col = content_len / colspan
                    for dc in range(colspan):
                        idx = origin_c + dc
                        if idx < num_cols:
                            cell_lengths.setdefault(idx, []).append(per_col)

        widths_cm = _calc_column_widths(cell_lengths, num_cols, config)
        for i, w in enumerate(widths_cm):
            table.columns[i].width = Cm(w)
    except Exception:
        pass


def _contains_symbol_chars(text):
    """检查文本是否包含 emoji 或特殊 Unicode 符号"""
    for c in text:
        cp = ord(c)
        if 0x2600 <= cp <= 0x27BF:  # Misc Symbols + Dingbats
            return True
        if 0x1F000 <= cp <= 0x1FFFF:  # Emoji
            return True
        if 0x2700 <= cp <= 0x27BF:  # Dingbats
            return True
    return False


def hex_to_rgb(hex_color: str):
    """将十六进制颜色转换为 RGBColor"""
    from docx.shared import RGBColor
    hex_color = hex_color.lstrip('#')
    if len(hex_color) == 6:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        return RGBColor(r, g, b)
    return RGBColor(0, 0, 0)  # 默认黑色
